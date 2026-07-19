"""Viewer reader plugin for Word documents (the ``documents`` extra).

Extracts plain text from ``.docx`` files (paragraphs, lists, tables)
for the built-in viewer. When tables are present, returns a table view
with proper formatting. Registers no extensions when the
``python-docx`` package (part of the ``documents`` extra, imported as
``docx``) isn't installed, per the guard convention documented in
``plugins/__init__.py``.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from linux_commander.vfs import FileSystem, VfsPath

if TYPE_CHECKING:
    from linux_commander.plugins import ViewDocument

try:
    import docx  # type: ignore[import-not-found]
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    VIEW_EXTENSIONS: tuple[str, ...] = (".docx",)
except ImportError:
    docx = None  # type: ignore[assignment]
    VIEW_EXTENSIONS = ()


def _iter_block_items(parent):
    """Yield paragraphs and tables in document order."""
    from docx.oxml.ns import qn

    for child in parent.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _paragraph_to_text(paragraph: Paragraph) -> str:
    """Convert a paragraph to plain text, preserving list formatting."""
    text = paragraph.text
    # Check for list numbering
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        # This is a list item - add a bullet prefix
        text = "• " + text
    return text


def _table_to_rows(table: Table) -> list[list[str]]:
    """Convert a docx table to a list of rows (list of cell texts)."""
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    return rows


def read_document(host_fs: FileSystem, path: VfsPath) -> ViewDocument:
    """Extract content from ``path`` as a document preview.

    Returns a table view if the document contains tables, otherwise text view.
    """
    from linux_commander.plugins import MAX_PREVIEW_ROWS, ViewDocument, cleanup_temp, materialize

    real_path = materialize(host_fs, path)
    try:
        document = docx.Document(str(real_path))

        # Check if document has tables
        has_tables = len(document.tables) > 0

        if has_tables:
            # Return table view - combine all tables
            all_rows: list[list[str]] = []
            truncated = False
            for table_idx, table in enumerate(document.tables):
                if table_idx > 0:
                    # Add separator row
                    num_cols = len(table.rows[0].cells)
                    all_rows.append([f"--- Table {table_idx + 1} ---"] + [""] * (num_cols - 1))
                table_rows = _table_to_rows(table)
                for _row_idx, row in enumerate(table_rows):
                    if len(all_rows) >= MAX_PREVIEW_ROWS:
                        truncated = True
                        break
                    all_rows.append(row)
                if truncated:
                    break

            return ViewDocument(kind="table", rows=all_rows, truncated=truncated, meta={})
        else:
            # Text view for documents without tables
            parts: list[str] = []
            for block in _iter_block_items(document):
                if isinstance(block, Paragraph):
                    parts.append(_paragraph_to_text(block))
                elif isinstance(block, Table):
                    table_rows = _table_to_rows(block)
                    for row in table_rows:
                        parts.append("\t".join(row))
            return ViewDocument(kind="text", text="\n".join(parts), meta={})
    finally:
        cleanup_temp(real_path)
