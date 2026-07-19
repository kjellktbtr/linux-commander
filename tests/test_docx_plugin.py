"""Tests for the docx viewer reader plugin (the ``documents`` extra).

Fixtures are built with python-docx directly, so these run anywhere the
optional dependency is installed. Skipped entirely when ``docx`` (the
``python-docx`` package) isn't installed.
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

import docx  # noqa: E402

from linux_commander.plugins import viewer_plugin_for_name  # noqa: E402
from linux_commander.plugins.docx_plugin import read_document  # noqa: E402
from linux_commander.vfs import LocalFileSystem, VfsPath  # noqa: E402

_FS = LocalFileSystem()


def _local_vpath(path: Path) -> VfsPath:
    return _FS.from_path(path)


def test_docx_extension_is_discovered() -> None:
    mod = viewer_plugin_for_name("letter.docx")
    assert mod is not None
    assert mod.read_document is read_document


def test_read_document_extracts_paragraph_text(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("Hello, World!")
    document.add_paragraph("Second paragraph.")
    path = tmp_path / "doc.docx"
    document.save(str(path))

    doc = read_document(_FS, _local_vpath(path))
    assert doc.kind == "text"
    assert "Hello, World!" in doc.text
    assert "Second paragraph." in doc.text


def test_read_document_extracts_table_cell_text(tmp_path: Path) -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "left"
    table.rows[0].cells[1].text = "right"
    path = tmp_path / "table.docx"
    document.save(str(path))

    doc = read_document(_FS, _local_vpath(path))
    assert doc.kind == "table"
    assert doc.rows == [["left", "right"]]


def test_read_document_empty_document_returns_empty_text(tmp_path: Path) -> None:
    document = docx.Document()
    path = tmp_path / "empty.docx"
    document.save(str(path))

    doc = read_document(_FS, _local_vpath(path))
    assert doc.kind == "text"
    assert doc.text == ""
